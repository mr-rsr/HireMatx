"""Resume processing service."""

import hashlib
import os
from datetime import datetime
from pathlib import Path

import structlog
from PyPDF2 import PdfReader
from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import get_settings
from app.models.resume import Resume, ResumeStatus
from app.models.user import User
from app.services.ai_service import AIService

logger = structlog.get_logger()


class ResumeService:
    """Service for resume processing."""

    def __init__(self, db: AsyncSession):
        self.db = db
        self.settings = get_settings()
        self.ai_service = AIService()

    async def upload_resume(
        self,
        user: User,
        filename: str,
        content: bytes,
        file_type: str,
    ) -> Resume:
        """Upload and store a resume."""
        # Calculate file hash for deduplication
        file_hash = hashlib.sha256(content).hexdigest()

        # Check for duplicate
        existing = await self.db.execute(
            select(Resume).where(
                Resume.user_id == user.id,
                Resume.file_hash == file_hash,
            )
        )
        if existing.scalar_one_or_none():
            raise ValueError("This resume has already been uploaded")

        # Create storage path
        storage_dir = Path("storage/resumes") / str(user.id)
        storage_dir.mkdir(parents=True, exist_ok=True)

        # Generate unique filename
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        stored_filename = f"{timestamp}_{filename}"
        file_path = storage_dir / stored_filename

        # Save file
        with open(file_path, "wb") as f:
            f.write(content)

        # Check if this is the first resume (make it primary)
        existing_count = await self.db.execute(
            select(Resume).where(Resume.user_id == user.id, Resume.is_active == True)
        )
        is_first = len(list(existing_count.scalars().all())) == 0

        # Create resume record
        resume = Resume(
            user_id=user.id,
            filename=filename,
            file_type=file_type,
            file_size=len(content),
            file_path=str(file_path),
            file_hash=file_hash,
            status=ResumeStatus.PENDING,
            is_primary=is_first,  # First resume is automatically primary
        )
        self.db.add(resume)
        await self.db.flush()

        return resume

    async def process_resume(self, resume: Resume) -> Resume:
        """Process and analyze a resume."""
        resume.status = ResumeStatus.PROCESSING
        await self.db.flush()

        try:
            # Extract text based on file type
            if resume.file_type.lower() == "pdf":
                text = self._extract_pdf_text(resume.file_path)
            elif resume.file_type.lower() in ["docx", "doc"]:
                text = self._extract_docx_text(resume.file_path)
            else:
                raise ValueError(f"Unsupported file type: {resume.file_type}")

            resume.raw_text = text

            # Analyze with AI
            analysis = await self.ai_service.analyze_resume(text)

            resume.ai_summary = analysis.get("summary")
            resume.ai_skills_extracted = {"skills": analysis.get("skills", [])}
            resume.ai_experience_level = analysis.get("experience_level")
            resume.ai_job_titles = {"titles": analysis.get("suggested_titles", [])}
            resume.parsed_data = {
                "strengths": analysis.get("strengths", []),
                "improvement_suggestions": analysis.get("improvement_suggestions", []),
                "industries": analysis.get("industries", []),
                "ats_score": analysis.get("ats_score"),
                "ats_suggestions": analysis.get("ats_suggestions", []),
            }

            resume.status = ResumeStatus.PROCESSED
            resume.processed_at = datetime.utcnow()

        except Exception as e:
            logger.error("resume_processing_error", error=str(e), resume_id=resume.id)
            resume.status = ResumeStatus.FAILED
            resume.error_message = str(e)

        await self.db.flush()
        return resume

    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n".join(text_parts)

    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        return "\n".join(text_parts)

    async def get_user_resumes(self, user: User) -> list[Resume]:
        """Get all resumes for a user."""
        result = await self.db.execute(
            select(Resume)
            .where(Resume.user_id == user.id, Resume.is_active == True)
            .order_by(Resume.created_at.desc())
        )
        return list(result.scalars().all())

    async def get_primary_resume(self, user: User) -> Resume | None:
        """Get user's primary resume."""
        result = await self.db.execute(
            select(Resume).where(
                Resume.user_id == user.id,
                Resume.is_primary == True,
                Resume.is_active == True,
            )
        )
        return result.scalar_one_or_none()

    async def set_primary_resume(self, user: User, resume_id: int) -> Resume | None:
        """Set a resume as primary."""
        # Clear existing primary
        result = await self.db.execute(
            select(Resume).where(
                Resume.user_id == user.id,
                Resume.is_primary == True,
            )
        )
        for resume in result.scalars():
            resume.is_primary = False

        # Set new primary
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user.id,
            )
        )
        resume = result.scalar_one_or_none()
        if resume:
            resume.is_primary = True
            await self.db.flush()

        return resume

    async def delete_resume(self, user: User, resume_id: int) -> bool:
        """Soft delete a resume."""
        result = await self.db.execute(
            select(Resume).where(
                Resume.id == resume_id,
                Resume.user_id == user.id,
            )
        )
        resume = result.scalar_one_or_none()

        if resume:
            resume.is_active = False
            await self.db.flush()
            return True

        return False

    async def get_resume_text(self, resume: Resume) -> str | None:
        """Get the extracted text from a resume."""
        if resume.raw_text:
            return resume.raw_text

        # Re-extract if not available
        if resume.status == ResumeStatus.PROCESSED:
            return None

        try:
            if resume.file_type.lower() == "pdf":
                return self._extract_pdf_text(resume.file_path)
            elif resume.file_type.lower() in ["docx", "doc"]:
                return self._extract_docx_text(resume.file_path)
        except Exception as e:
            logger.error("resume_text_extraction_error", error=str(e))

        return None
